import io
import sys
import os
from google.cloud import vision
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tkinter as tk
from tkinter import filedialog, messagebox, ttk


# Check if we're running as a script or a frozen exe
if getattr(sys, 'frozen', False):
    application_path = sys._MEIPASS  # If we're an exe, this is the path of the exe file
else:
    application_path = os.path.dirname(os.path.abspath(__file__))  # If we're a script, this is the path of the script file
    
# Path to the Google API credentials
google_json_path = os.path.join(application_path, 'google.json')

# Set environment variable for Google API credentials
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = google_json_path

def process_image(image_path, progress_var, window, progress_bar):
    with io.open(image_path, 'rb') as image_file:
        content = image_file.read()

    client = vision.ImageAnnotatorClient()
    image = vision.Image(content=content)
    response = client.document_text_detection(image=image)

    document = Document()

    total_paragraphs = sum(1 for _ in response.full_text_annotation.pages)
    progress_var.set(0)
    progress_bar["maximum"] = total_paragraphs  # Set the maximum value of the progress bar

    for page in response.full_text_annotation.pages:
        for block in page.blocks:
            for paragraph in block.paragraphs:
                word_text = " ".join(["".join([symbol.text for symbol in word.symbols]) for word in paragraph.words])
                p = document.add_paragraph(word_text + " ")
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT 
                p.style = document.styles['Normal']
                font = p.style.font
                font.rtl = True  
                progress_var.set(progress_var.get() + 1)
                window.update_idletasks()

    return document

def create_gui():
    window = tk.Tk()
    window.title("Image to Text Converter")
    window.geometry("500x200")

    progress_var = tk.DoubleVar()
    progress_bar = ttk.Progressbar(window, length=200, variable=progress_var)
    progress_bar.grid(row=3, column=1, padx=10, pady=10)

    filename_entry = tk.Entry(window, width=40)  # Define filename_entry before using it in the button command
    filename_entry.grid(row=1, column=1, padx=10)

    image_entry = tk.Entry(window, width=40)
    image_entry.grid(row=0, column=1, padx=10)

    def select_image():
        file_path = filedialog.askopenfilename(filetypes=[("Image Files", "*.jpg;*.jpeg;*.png")])
        image_entry.delete(0, tk.END) 
        image_entry.insert(0, file_path) 

    def save_docx():
        output_path = filename_entry.get()
        if output_path:
            result = process_image(image_entry.get(), progress_var, window, progress_bar)  
            result.save(output_path)
            messagebox.showinfo("Process Finished", "The document has been saved successfully.")

    tk.Label(window, text="Select Image:").grid(row=0, column=0, padx=10, pady=10)
    tk.Button(window, text="Browse", command=select_image).grid(row=0, column=2, padx=10)

    tk.Label(window, text="Save As:").grid(row=1, column=0, padx=10, pady=10)
    tk.Button(window, text="Browse", command=lambda: filename_entry.insert(tk.END, filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")]))).grid(row=1, column=2, padx=10)

    tk.Button(window, text="Convert and Save", command=save_docx).grid(row=2, column=1, padx=10, pady=10)

    window.mainloop()

# Start the GUI  
if __name__ == "__main__":
    create_gui()